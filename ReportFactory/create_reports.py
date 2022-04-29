import plotly.express as px
import plotly.graph_objects as go
import re
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Cm
from datetime import datetime
from Utilities.create_sub_dfs import return_frames



def create_worddoc(var_dict, baseline_dict, df_CapexBudgetDemandCatalogue2, df_intersects2, df_EntireSet, df_perward,
                   df_CapexBudgetDemandCatalogue3):
    """
    This module creates a spatial report in a MSWord File
    """
    # *********************************************************************************************************
    #                                INITIALISATION OF KEY VARIABLES HAPPENS HERE                             #
    # *********************************************************************************************************
    global teller, fig_nr, fig, tbl_nr, tbl, heading_no
    fig = {}
    tbl ={}
    heading_no = 1
    fig_nr = 1
    tbl_nr = 1
    # *********************************************************************************************************
    #                                      END OF INITIALISE KEY VARIABLES                                    #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #                                       ALL SUB-FUNCTIONS GO HERE                                         #
    # *********************************************************************************************************
    def par_formatter(paragraphs):
        """
        This function within a function formats paragraphs using normal, bold or italic.
        """
        paragraph = document.add_paragraph()
        for key, value in paragraphs.items():
            if " n" in key:  # Previous formatting
                paragraph.add_run(value)
            elif "al" in key: # Normal
                document.add_paragraph(value, style="Normal")
            elif " b" in key:  # Bold
                paragraph.add_run(value).bold = True
            elif " i" in key:  # Italic
                paragraph.add_run(value).italic = True
            elif " u" in key:  # Underline
                paragraph.add_run(value).underline = True
            elif "em" in key:  # Style: Emphasis
                paragraph.add_run(value).style = 'Emphasis'
            elif "tl" in key:  # Style: Title
                document.add_paragraph(value, style="Title")
            elif "st" in key:  # Style: SubTitle
                try:
                    document.add_paragraph(value, style="Subtitle")
                except KeyError:
                    document.add_paragraph(value, style="Strong")
            elif 'se' in key:
                document.add_paragraph(value, style="Section")
            elif "lb" in key:  # Style: List Bullet
                try:
                    document.add_paragraph(value, style="List Bullet")
                except KeyError:
                    document.add_paragraph(value, style="Bulleted list")
            elif "xx" in key:
                document.add_paragraph(value, style="Bulleted list")
            elif "cp" in key:  # Style: Caption
                document.add_paragraph(value, style="Caption")


    def head_formatter(headings):
        """
         This function within a function formats headings using the heading levels as input.
        """
        for key, value in headings.items():
            if key == "0":
                    document.add_heading(value, level=0)
            elif key == "1":
                    document.add_heading(value, level=1)
            elif key == "2":
                    document.add_heading(value, level=2)
            elif key == "3":
                    document.add_heading(value, level=3)
            elif key == "4":
                    document.add_heading(value, level=4)


    def format_budget(n):
        """
        This function formats a number entry into a currency by putting an 'R' in front and delimiting with commas
        """
        currency = "R{:0,.0f}".format(n).replace('R-', '-R')
        # ${:0,.0f}
        return currency


    def format_percent(n):
        """
        This function formats a number entry into a percentage by putting an '%' in the back and delimiting with 2
        decimals
        """
        n = float(n)
        n = n * 100
        percentage = "{:.2f}%".format(n)
        return percentage


    def build_bar_plots_1(number_of_plots, fig_nr, df_subs):
        """
        This function automates the building of bar plots based on the number of sub_plots required as provided
        by the parameter 'number_of_plots' which in turn gets created by len(df_subs) under the 'create_sub_dfs'
        module.
        """
        fig_t = {}
        temp_fig = {}

        for plot_no in range(1, number_of_plots + 1):
            if number_of_plots > 1:
                temp_fig[plot_no] = f"{fig_nr}.{plot_no}"
            else:
                temp_fig[plot_no] = f"{fig_nr}"
            fig_t[plot_no] = {}
            fig_t[plot_no][temp_fig[plot_no]] = px.bar(df_subs[plot_no - 1], x=f'Projects per {SpatialFeatureChoice}',
                                            y=SpatialFeatureChoice,
                                            color='Capital All Years',
                                            height=1500, width=1100, orientation='h',
                                                       color_continuous_scale=px.colors.sequential.Rainbow)
            # color_continuous_scale=px.colors.sequential.Burg
            # Sort images to follow each other
            fig_t[plot_no][temp_fig[plot_no]].update_yaxes(categoryorder='total descending')
            fig_t[plot_no][temp_fig[plot_no]].update_coloraxes(colorbar_tickprefix='R', colorbar_tickformat=',.')
            # Write images to png
            fig_t[plot_no][temp_fig[plot_no]].write_image(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png")
            # Add a page break
            # document.add_page_break()
            temp_paragraphs = {}
            if number_of_plots > 1:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Projects and Capital Demand per " \
                                       f"{SpatialFeatureChoice_Text_Single} ({plot_no}/{number_of_plots})"
            else:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Projects and Capital Demand per " \
                                       f"{SpatialFeatureChoice_Text_Single}"
            par_formatter(temp_paragraphs)
            document.add_picture(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png", width=Cm(17))


    def build_bar_plots_2(number_of_plots, fig_nr, df_subs):
        """
        This function automates the building of bar plots based on the number of sub_plots required as provided
        by the parameter 'number_of_plots' which in turn gets created by len(df_subs) under the 'create_sub_dfs'
        module.
        """
        fig_t = {}
        temp_fig = {}

        for plot_no in range(1, number_of_plots + 1):
            if number_of_plots > 1:
                temp_fig[plot_no] = f"{fig_nr}.{plot_no}"
            else:
                temp_fig[plot_no] = f"{fig_nr}"
            fig_t[plot_no] = {}
            fig_t[plot_no][temp_fig[plot_no]] = px.histogram(df_subs[plot_no - 1], y=SpatialFeatureChoice,
                                                             x=column_name_list,
                                                             nbins=len(df_subs[plot_no - 1]), height=1500, width=1100,
                                                             orientation='h', color_discrete_sequence=colors)
            fig_t[plot_no][temp_fig[plot_no]].update_yaxes(categoryorder='total descending')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(xaxis_tickprefix='R', xaxis_tickformat=',.')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(barmode='stack')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(legend_title="Capital Demand MTREF")

            # Write images to png
            fig_t[plot_no][temp_fig[plot_no]].write_image(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png")
            # Add a page break
            # document.add_page_break()
            temp_paragraphs = {}
            if number_of_plots > 1:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Capital Demand per {SpatialFeatureChoice_Text_Single} - " \
                                       f"All Years ({plot_no}/{number_of_plots})"
            else:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Capital Demand per {SpatialFeatureChoice_Text_Single} - " \
                                       f"All Years"
            par_formatter(temp_paragraphs)
            document.add_picture(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png", width=Cm(17))


    def build_bar_plots_3(number_of_plots, fig_nr, df_subs):
        """
        This function automates the building of bar plots based on the number of sub_plots required as provided
        by the parameter 'number_of_plots' which in turn gets created by len(df_subs) under the 'create_sub_dfs'
        module.
        """
        fig_t = {}
        temp_fig = {}

        for plot_no in range(1, number_of_plots + 1):
            if number_of_plots > 1:
                temp_fig[plot_no] = f"{fig_nr}.{plot_no}"
            else:
                temp_fig[plot_no] = f"{fig_nr}"
            fig_t[plot_no] = {}
            fig_t[plot_no][temp_fig[plot_no]] = px.histogram(df_subs[plot_no - 1], y=SpatialFeatureChoice,
                                                             x=column_name_list[0:3],
                                                             nbins=len(df_subs[plot_no - 1]), height=1500, width=1100,
                                                             orientation='h', color_discrete_sequence=colors)
            fig_t[plot_no][temp_fig[plot_no]].update_yaxes(categoryorder='total descending')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(xaxis_tickprefix='R', xaxis_tickformat=',.')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(barmode='stack')
            fig_t[plot_no][temp_fig[plot_no]].update_layout(legend_title="Capital Demand All Years")
            # Write images to png
            fig_t[plot_no][temp_fig[plot_no]].write_image(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png")
            # Add a page break
            # document.add_page_break()
            temp_paragraphs = {}
            if number_of_plots > 1:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Capital Demand per {SpatialFeatureChoice_Text_Single} - " \
                                       f"MTREF Period ({plot_no}/{number_of_plots})"
            else:
                temp_paragraphs[
                    f'{plot_no} cp'] = f"Figure {heading_no}.{temp_fig[plot_no]}: Capital Demand per {SpatialFeatureChoice_Text_Single} - " \
                                       f"MTREF Period"
            par_formatter(temp_paragraphs)
            document.add_picture(f"./static/images/fig{temp_fig[plot_no]}{plot_no}.png", width=Cm(17))
    # *********************************************************************************************************
    #                                       END OF SUB-FUNCTIONS                                              #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #                                       ALL VARIABLES GO HERE                                             #
    # *********************************************************************************************************
    colors = ['blue', 'green', 'red', 'orange', 'yellow', 'beige', 'teal', 'purple', 'gray']
    # Unpack the variables that were passed to the function in a dictionary called var_dict
    url_choice = var_dict['url_choice']
    entity_choice = var_dict['entity_choice']
    SpatialFeatureChoice = var_dict['SpatialFeatureChoice']
    Layer_List = var_dict['Layer_List']
    total_datapoints = var_dict['total_datapoints']
    intersecting = var_dict['intersecting']
    no_intersects = var_dict['no_intersects']
    perc_no_intersects = (no_intersects / total_datapoints * 100)
    chosen_feature_qty = var_dict['chosen_feature_qty']
    column_name_list = var_dict['column_name_list']
    list_of_years = var_dict['list_of_years']
    list_of_features = var_dict['list_of_features']
    SpatialFeatureChoice_Text = SpatialFeatureChoice.replace("_", "")
    if SpatialFeatureChoice_Text[-1].lower() == 's':
        SpatialFeatureChoice_Text_Single = SpatialFeatureChoice_Text[:-1]
    else:
        SpatialFeatureChoice_Text_Single = SpatialFeatureChoice_Text
    # Set up other vars
    timenow = datetime.today().strftime('%Y-%m-%d-%H_%M')  # Stamp the date and time
    datenow = datetime.today().strftime('%Y-%m-%d')  # Stamp the date
    urlname = url_choice.split(sep=".", maxsplit=10)[1].capitalize()
    word_file_name = f"{urlname}_SpatialRep_{timenow}.docx"
    df_EntireSet_Temp = df_EntireSet
    # Create variables about the projects
    maximum_projects = df_EntireSet[f'Projects per {SpatialFeatureChoice}'].max()
    maximum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {SpatialFeatureChoice}'] == df_EntireSet[
        f'Projects per {SpatialFeatureChoice}'].max(), f'{SpatialFeatureChoice}'].iloc[0]
    minimum_projects = df_EntireSet[f'Projects per {SpatialFeatureChoice}'].min()
    minimum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {SpatialFeatureChoice}'] == df_EntireSet[
        f'Projects per {SpatialFeatureChoice}'].min(), f'{SpatialFeatureChoice}'].iloc[0]
    average_projects = df_EntireSet[f'Projects per {SpatialFeatureChoice}'].mean()
    seventy_fifth_projects = df_EntireSet[f'Projects per {SpatialFeatureChoice}'].quantile(q=0.75)
    sum_projects = df_EntireSet[f'Projects per {SpatialFeatureChoice}'].sum()
    min_projects_perc_of_total = format_percent(minimum_projects / sum_projects)
    max_projects_perc_of_total = format_percent(maximum_projects / sum_projects)
    # Re-sort the df to get the top 5 wards with the most projects
    df_EntireSet_Temp.sort_values(f'Projects per {SpatialFeatureChoice}', inplace=True, ascending=True)
    sum_top_five_projects = df_EntireSet_Temp[f'Projects per {SpatialFeatureChoice}'].tail().sum()
    sum_top_five_projects_perc_of_total = format_percent(sum_top_five_projects / sum_projects)
    top_five_projfeat_text = " and".join(
        ", ".join(df_EntireSet_Temp[f'{SpatialFeatureChoice}'].tail().tolist()).rsplit(',', 1))
    # Create variables about the costs
    maximum_cost_float = df_EntireSet['Capital All Years'].max()
    maximum_cost = format_budget(maximum_cost_float)
    maximum_cost_feature = \
        df_EntireSet.loc[df_EntireSet['Capital All Years'] == df_EntireSet[
            'Capital All Years'].max(), f'{SpatialFeatureChoice}'].iloc[0]
    minimum_cost_float = df_EntireSet['Capital All Years'].min()
    minimum_cost = format_budget(minimum_cost_float)
    minimum_cost_feature = \
        df_EntireSet.loc[df_EntireSet['Capital All Years'] == df_EntireSet[
            'Capital All Years'].min(), f'{SpatialFeatureChoice}'].iloc[
            0]
    average_cost_float = df_EntireSet['Capital All Years'].mean()
    average_cost = format_budget(average_cost_float)
    seventy_fifth_cost_float = df_EntireSet['Capital All Years'].quantile(q=0.75)
    seventy_fifth_cost = format_budget(seventy_fifth_cost_float)
    sum_cost_float = df_EntireSet['Capital All Years'].sum()
    sum_cost = format_budget(sum_cost_float)
    max_cost_perc_of_total_float = df_EntireSet['Capital All Years'].max() / df_EntireSet['Capital All Years'].sum()
    min_cost_perc_of_total_float = df_EntireSet['Capital All Years'].min() / df_EntireSet['Capital All Years'].sum()
    max_cost_perc_of_total = format_percent(max_cost_perc_of_total_float)
    min_cost_perc_of_total = format_percent(min_cost_perc_of_total_float)
    # Re-sort the df to get the top 5 capital demand
    df_EntireSet_Temp.sort_values('Capital All Years', inplace=True, ascending=True)
    sum_top_five_cost_float = df_EntireSet_Temp['Capital All Years'].tail().sum()
    sum_top_five_cost = format_budget(sum_top_five_cost_float)
    sum_top_five_cost_perc_of_total_float = sum_top_five_cost_float / sum_cost_float
    sum_top_five_cost_perc_of_total = format_percent(sum_top_five_cost_perc_of_total_float)
    top_five_capfeat_text = " and".join(
        ", ".join(df_EntireSet_Temp[f'{SpatialFeatureChoice}'].tail().tolist()).rsplit(',', 1))
    # Create variables for the table summary
    maximum_capital_float_MTREF = df_EntireSet['Capital MTREF'].max()
    maximum_capital_MTREF = format_budget(maximum_capital_float_MTREF)
    maximum_capital_feature_MTREF = df_EntireSet.loc[df_EntireSet['Capital MTREF'] ==
                                                     df_EntireSet['Capital MTREF'].max(),
                                                     f'{SpatialFeatureChoice}'].iloc[0]
    minimum_capital_float_MTREF = df_EntireSet['Capital MTREF'].min()
    minimum_capital_MTREF = format_budget(minimum_capital_float_MTREF)
    minimum_capital_feature_MTREF = df_EntireSet.loc[df_EntireSet['Capital MTREF'] ==
                                                     df_EntireSet['Capital MTREF'].min(),
                                                     f'{SpatialFeatureChoice}'].iloc[0]
    average_cost_MTREF_float = df_EntireSet['Capital MTREF'].mean()
    average_cost_MTREF = format_budget(average_cost_MTREF_float)
    seventy_fifth_cost_float_MTREF = df_EntireSet['Capital MTREF'].quantile(q=0.75)
    seventy_fifth_cost_MTREF = format_budget(seventy_fifth_cost_float_MTREF)
    sum_cost_float_MTREF = df_EntireSet['Capital MTREF'].sum()
    sum_cost_MTREF = format_budget(sum_cost_float_MTREF)
    max_cost_perc_of_total_float_MTREF = df_EntireSet['Capital MTREF'].max() / df_EntireSet['Capital MTREF'].sum()
    min_cost_perc_of_total_float_MTREF = df_EntireSet['Capital MTREF'].min() / df_EntireSet['Capital MTREF'].sum()
    max_cost_perc_of_total_MTREF = format_percent(max_cost_perc_of_total_float_MTREF)
    min_cost_perc_of_total_MTREF = format_percent(min_cost_perc_of_total_float_MTREF)
    list_of_years_str = []
    for year in list_of_years:
        list_of_years_str.append(str(year))
    mtref_years_text = " and".join(", ".join(list_of_years_str[0:3]).rsplit(',', 1))
    # *********************************************************************************************************
    #                                       END OF VARIABLES AREA                                             #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #                                       ALL REPORT CONTENT GOES HERE                                      #
    # *********************************************************************************************************
    def create_content(heading_no, fig_nr, tbl_nr):
        """
        This function contains all the report conntent is dictionaries. The dictionaries are passed into a
        format interpreter (seperate function) that generates the content in the actual docx file.
        """
        if heading_no == 0:  # General Background Info
            paragraphs0 = {"1 n": "This document was created ",
                           "2 i": f"{timenow} ",
                           "3 n": "from the ",
                           "4 b": f"{entity_choice} ",
                           "5 n": "CP3 system."}
            paragraphs1 = {"1 n": "The baseline ",
                           "2 b": f"'{baseline_dict['Name']}'",
                           "3 n": " with description: ",
                           "4 b": f"'{baseline_dict['Description']}'",
                           "5 n": " was used. ",
                           "6 n": "The spatial feature that was selected for the purpose of this report was: ",
                           "7 b": f"'{SpatialFeatureChoice_Text}'",
                           "8 n": "."}
            paragraphs2 = {"1 n": f"This is a spatial feature report for the ",
                           "2 i": f"{entity_choice}.",
                           "3 n": f" The data contained in this report was sourced directly from the municipality's "
                                  f"live CP3 system. The information in this report therefore reflects the data as it"
                                  f" was on ",
                           "4 i": f"{datenow}",
                           "5 n": f" when this report was requested. Any subsequent updates to the data contained the"
                                  f" CP3 System related to the applicable baseline from which this reports was drawn, "
                                  f"would therefore not reflect in this report. \n\nFor the ",
                           "6 i": f"{entity_choice}",
                           "7 i": f", the following spatial features are available for the purpose of developing ",
                           "8 u": f"'Spatial Query Reports'",
                           "9 n": f" (similar to this report):"}
            teller = 10
            for feature in Layer_List:
                paragraphs2[f"{teller} xx"] = feature.replace("_","")
                teller += 1
            paragraphs3 = {"1 n": "We hope you find this useful! Sincerely, ",
                           "2 i": "The Novus3 Team."}
            # *********************************************************************************************************
            #                                       Chapter 0 - General Background Info                               #
            # *********************************************************************************************************
            par_formatter(paragraphs0)
            par_formatter(paragraphs1)
            par_formatter(paragraphs2)
            document.add_picture(f"./static/images/CP3logo.png", width=Cm(3))
            par_formatter(paragraphs3)
            document.add_page_break()
            # *********************************************************************************************************
            #                                       END Chapter 0 - General Background Info                           #
            # *********************************************************************************************************
        elif heading_no == 1:  # Introduction
            headings1 = {"1": "Introduction"}
            # Put together the dictionary
            paragraphs4 = {"1 n": f"The baseline that was queried for this report contains ",
                           "2 i": f"{total_datapoints}",
                           "3 n": f" projects of which a total of ",
                           "4 i": f"{no_intersects} ({'{0:.3g}'.format(perc_no_intersects)}%)",
                           "5 n": f" projects do not have any recorded intersection (overlap) with "
                                  f"{SpatialFeatureChoice_Text}. The following spatial features are available for "
                                  f"queries using this API profile:"
                           }
            teller = 6
            for feature in Layer_List:
                paragraphs4[f"{teller} xx"] = feature.replace("_","")
                teller += 1
            paragraphs4[f"{teller} al"] = f"\nThe probable reasons for the {no_intersects} non intersecting projects" \
                                          f" reported are that:"
            teller += 1
            paragraphs4[f"{teller} xx"] = f"these projects simply do not have a location associated with them yet " \
                                          f"(most frequently the case with 'no-intersect' projects) or;"
            teller += 1
            paragraphs4[f"{teller} xx"] = f"the location that was captured for some of these projects, may have been" \
                                          f" in the wrong place (e.g. outside the boundaries of {urlname} - this is " \
                                          f"very rarely the reason)."
            teller += 1
            paragraphs4[f"{teller} al"] = f"\nIt is important to take note of the " \
                                          f"{'{0:.3g}'.format(perc_no_intersects)}% projects that do not intersect " \
                                          f"with {SpatialFeatureChoice_Text} feature when appraising this report."
            if perc_no_intersects <= 5:
                teller += 1
                paragraphs4[f"{teller} al"] = f"The {'{0:.3g}'.format(perc_no_intersects)}% of non-intersecting " \
                                              f"projects is a relative low percentage and therefore not much cause of" \
                                              f" concern for this particular report."
            teller += 1
            paragraphs4[f"{teller} cp"] = f"Figure {heading_no}.{fig_nr}: Intersecting vs Non-Intersecting Projects"
            # *********************************************************************************************************
            #                                       1. Introduction                                                   #
            # *********************************************************************************************************
            head_formatter(headings1)
            par_formatter(paragraphs4)
            # *********************************************************************************************************
            #                                       Add Figure 1.1                                                    #
            # *********************************************************************************************************
            fig[fig_nr] = px.pie(df_intersects2, values='Projects', names='Intersects')
            fig[fig_nr].update_traces(marker=dict(colors=colors))
            fig[fig_nr].write_image(f"./static/images/fig{fig_nr}.png")
            document.add_picture(f"./static/images/fig{fig_nr}.png", width=Cm(15))
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 1.1                                                    #
            # *********************************************************************************************************
            # *********************************************************************************************************
            #                                       END 1. Introduction                                               #
            # *********************************************************************************************************
        elif heading_no == 2:  # Total no of projs and capital demand
            headings2 = {"1": f"Total Number of Projects and Total Capital Demand per "
                              f"{SpatialFeatureChoice_Text_Single} - All Years"}
            paragraphs5a = {"1 n": f"Table {heading_no}.{tbl_nr} below provides a concise summary of key statistical"
                                   f" insights regarding the number of projects that are requesting funding in each "
                                   f"of the {SpatialFeatureChoice_Text}.",
                            "2 cp": f"Table {heading_no}.{tbl_nr}: {SpatialFeatureChoice_Text} Analysis - Number of"
                                    f" Projects"
                            }
            paragraphs5b = {"1 n": f"\nTable {heading_no}.{tbl_nr} below provides a concise summary of key statistical"
                                   f" insights regarding the total capital demand per "
                                   f"{SpatialFeatureChoice_Text_Single}.",
                            "2 cp": f"Table {heading_no}.{tbl_nr}: {SpatialFeatureChoice_Text} Analysis - Capital "
                                    f"Demand (R)"
                            }
            paragraphs5c = {"1 n": f"\nThe number of projects per {SpatialFeatureChoice_Text_Single} and the capital "
                                   f"demand per {SpatialFeatureChoice_Text_Single} needs to be looked at together for"
                                   f" better insight. Figures {heading_no}.{fig_nr} and {heading_no}.{fig_nr+1} assist"
                                   f" with further graphical insight into some of the numbers that are provided in "
                                   f"Tables {heading_no}.{tbl_nr} and {heading_no}.{tbl_nr+1}.",
                            "2 cp": f"Figure {heading_no}.{fig_nr}: The single {SpatialFeatureChoice_Text_Single} with"
                                    f" the highest value (Nr of Projects & Capital Demand) vs the rest"
                            }
            paragraphs6 = {"1 n": f"The highest number of projects in a single {SpatialFeatureChoice_Text_Single} is"
                                  f" in ",
                           "2 b": f"{maximum_projects_feature}. ",
                           "3 i": f"That is {maximum_projects} projects and this is {max_projects_perc_of_total} of "
                                  f"the total number of projects.",
                           "4 n": f"The highest capital demand in a single {SpatialFeatureChoice_Text_Single} is in ",
                           "5 b": f"{maximum_cost_feature}. ",
                           "6 i": f"That is {maximum_cost} and this amounts to {max_cost_perc_of_total} of the total"
                                  f" capital demand."}
            paragraphs5d = {"1 cp": f"Figure {heading_no}.{fig_nr+1}: The 5 {SpatialFeatureChoice_Text} with the highest"
                                    f" values (Nr of Projects & Capital Demand) vs the rest"
                            }
            paragraphs7 = {"1 n": f"The 5 {SpatialFeatureChoice_Text} that collectively have with the highest number"
                                  f" of projects are in ",
                           "2 b": f"{top_five_projfeat_text}. ",
                           "3 i": f"That is {sum_top_five_projects} projects and this is"
                                  f" {sum_top_five_projects_perc_of_total} of the total number of projects.",
                           "4 n": f"The 5 {SpatialFeatureChoice_Text} that collectively have the highest capital"
                                  f" demand are in ",
                           "5 b": f"{top_five_capfeat_text}. ",
                           "6 i": f"That is {sum_top_five_cost_perc_of_total} and this amounts to {sum_top_five_cost}"
                                  f" of the total capital demand."
                           }
            paragraphs8 = {
                "1 n": f"The following figures provide a bar graph representation of the number of projects per "
                       f"{SpatialFeatureChoice_Text_Single} with the colour of the bars representing the level of"
                       f" capital demand in each of these {SpatialFeatureChoice_Text}."}
            # *********************************************************************************************************
            #  2.	Total Number of Projects and Total Capital Demand per Spatial Feature (All Years)                 #
            # *********************************************************************************************************
            head_formatter(headings2)
            par_formatter(paragraphs5a)
            # *********************************************************************************************************
            #                                       Add Table 2.1                                                     #
            # *********************************************************************************************************
            table = document.add_table(rows=1, cols=4, style='List Table 4')
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Description'
            heading_cells[1].text = 'Info (1)'
            heading_cells[2].text = 'Info (2)'
            heading_cells[3].text = 'Info (3)'
            cells = table.add_row().cells
            cells[0].text = f"Number of projects in all {SpatialFeatureChoice_Text} - Overview"
            cells[1].text = f"Total in all {SpatialFeatureChoice_Text}:\n{sum_projects} projects"
            cells[2].text = f"Average per {SpatialFeatureChoice_Text}:\n{'{0:.3g}'.format(average_projects)} projects"
            cells[
                3].text = f"75th percentile of projects across all {SpatialFeatureChoice_Text}:" \
                          f"\n{seventy_fifth_projects} projects"
            cells = table.add_row().cells
            cells[0].text = f"The highest number of projects"
            cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {maximum_projects_feature}"
            cells[2].text = f"Number of projects in {maximum_projects_feature}:\n{maximum_projects}"
            cells[3].text = f"Percentage of total in {maximum_projects_feature}:\n{max_projects_perc_of_total}"
            cells = table.add_row().cells
            cells[0].text = f"The lowest number of projects:"
            cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {minimum_projects_feature}"
            cells[2].text = f"Number of projects in {minimum_projects_feature}:\n{minimum_projects}"
            cells[3].text = f"Percentage of total in {minimum_projects_feature}:\n{min_projects_perc_of_total}"
            if chosen_feature_qty > 10:  # If there are enough features to show 5 highest
                cells = table.add_row().cells
                cells[
                    0].text = f"The 5 {SpatialFeatureChoice_Text} that collectively have with the highest number" \
                              f" of projects"
                cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {top_five_projfeat_text}"
                cells[2].text = f"Number of projects in {top_five_projfeat_text} together:\n{sum_top_five_projects}"
                cells[
                    3].text = f"Percentage of total in {top_five_projfeat_text} together:" \
                              f"\n{sum_top_five_projects_perc_of_total}"
            # *********************************************************************************************************
            #                                       End Table 2.1                                                     #
            # *********************************************************************************************************
            tbl_nr += 1
            par_formatter(paragraphs5b)
            # *********************************************************************************************************
            #                                       Add Table 2.2                                                     #
            # *********************************************************************************************************
            table = document.add_table(rows=1, cols=4, style='List Table 4')
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Description'
            heading_cells[1].text = 'Info (1)'
            heading_cells[2].text = 'Info (2)'
            heading_cells[3].text = 'Info (3)'
            cells = table.add_row().cells
            cells[0].text = f"Capital demand in all {SpatialFeatureChoice_Text} - Overview"
            cells[1].text = f"Total capital demand in all {SpatialFeatureChoice_Text}:\n{sum_cost}"
            cells[2].text = f"The average capital demand per {SpatialFeatureChoice_Text}:\n{average_cost}"
            cells[
                3].text = f"The 75th percentile of capital demand for {SpatialFeatureChoice_Text}:" \
                          f"\n{seventy_fifth_cost}"
            cells = table.add_row().cells
            cells[0].text = f"The highest capital demand"
            cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {maximum_cost_feature}"
            cells[2].text = f"Capital Demand in {maximum_cost_feature}:\n{maximum_cost}"
            cells[3].text = f"Percentage of total in {maximum_cost_feature}:\n{max_cost_perc_of_total}"
            cells = table.add_row().cells
            cells[0].text = f"The lowest capital demand"
            cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {minimum_cost_feature}"
            cells[2].text = f"Capital Demand in {minimum_cost_feature}:\n{minimum_cost}"
            cells[3].text = f"Percentage of total in {minimum_cost_feature}:\n{min_cost_perc_of_total}"
            if chosen_feature_qty > 10:
                cells = table.add_row().cells
                cells[
                    0].text = f"The 5 {SpatialFeatureChoice_Text} that collectively have collective the highest " \
                              f"capital demand:"
                cells[1].text = f"{SpatialFeatureChoice_Text}:\nIn {top_five_capfeat_text}"
                cells[2].text = f"Capital Demand in {top_five_capfeat_text} together:\n{sum_top_five_cost}"
                cells[
                    3].text = f"Percentage of total in {top_five_capfeat_text} together:" \
                              f"\n{sum_top_five_cost_perc_of_total}"
            # Reset tbl_nr to 1 for next chapter
            tbl_nr += 1
            # *********************************************************************************************************
            #                                       End Table 2.2                                                     #
            # *********************************************************************************************************
            par_formatter(paragraphs5c)
            # *********************************************************************************************************
            #                                       Add Figure 2.1                                                    #
            # *********************************************************************************************************
            labels = ["Highest", "The Rest"]
            # Create subplots: use 'domain' type for Pie subplot
            fig[fig_nr] = make_subplots(rows=1, cols=2, specs=[[{'type': 'domain'}, {'type': 'domain'}]])
            fig[fig_nr].add_trace(go.Pie(labels=labels, values=[maximum_projects, (sum_projects - maximum_projects)],
                                         name="Nr of Projs"), 1, 1)
            fig[fig_nr].add_trace(
                go.Pie(labels=labels, values=[maximum_cost_float, (sum_cost_float - maximum_cost_float)],
                       name="Cap Dem (R)"), 1, 2)
            # Use `hole` to create a donut-like pie chart
            fig[fig_nr].update_traces(hole=.4)
            fig[fig_nr].update_traces(marker=dict(colors=colors))
            fig[fig_nr].update_layout(
                annotations=[dict(text='Nr of Projs', x=0.16, y=0.5, font_size=10, showarrow=False),
                             dict(text='Cap Dem (R)', x=0.85, y=0.5, font_size=10, showarrow=False)])
            fig[fig_nr].update_layout(legend=dict(orientation="h", y=0.99, x=0.35))

            fig[fig_nr].write_image(f"./static/images/fig{fig_nr}.1.png")
            document.add_picture(f"./static/images/fig{fig_nr}.1.png", width=Cm(17))
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 2.1                                                    #
            # *********************************************************************************************************
            par_formatter(paragraphs6)
            if chosen_feature_qty > 10:
                par_formatter(paragraphs5d)
                # ****************************************************************************************************
                #                                       Add Figure 2.2                                               #
                # ****************************************************************************************************
                # Create subplots: use 'domain' type for Pie subplot
                labels = ["5 Highest", "The Rest"]
                fig[fig_nr] = make_subplots(rows=1, cols=2, specs=[[{'type': 'domain'}, {'type': 'domain'}]])
                fig[fig_nr].add_trace(
                    go.Pie(labels=labels, values=[sum_top_five_projects, (sum_projects - sum_top_five_projects)],
                           name="Nr of Projs"), row=1, col=1)
                fig[fig_nr].add_trace(
                    go.Pie(labels=labels, values=[sum_top_five_cost_float, (sum_cost_float - sum_top_five_cost_float)],
                           name="Cap Dem (R)"), row=1, col=2)
                # Use `hole` to create a donut-like pie chart
                fig[fig_nr].update_traces(hole=.4)
                fig[fig_nr].update_traces(marker=dict(colors=colors))
                fig[fig_nr].update_layout(
                    annotations=[dict(text='Nr of Projs', x=0.16, y=0.5, font_size=10, showarrow=False),
                                 dict(text='Cap Dem (R)', x=0.85, y=0.5, font_size=10, showarrow=False)])
                fig[fig_nr].update_layout(legend=dict(orientation="h", y=0.99, x=0.35))

                fig[fig_nr].write_image(f"./static/images/fig{fig_nr}.2.png")
                document.add_picture(f"./static/images/fig{fig_nr}.2.png", width=Cm(17))
                fig_nr += 1
                # *********************************************************************************************************
                #                                       End Figure 2.2                                                    #
                # *********************************************************************************************************
                par_formatter(paragraphs7)
            par_formatter(paragraphs8)
            # *********************************************************************************************************
            #                                       Add Figure 2.3                                                    #
            # *********************************************************************************************************
            # Split and sort the main data frame
            df_subs = return_frames(df_Master=df_EntireSet, feature_qty=chosen_feature_qty, block_limit=40,
                                    sort_column=f'Projects per {SpatialFeatureChoice}')
            number_of_plots = len(df_subs)
            # Make the bar plots with number of projects and capital demand per project
            build_bar_plots_1(number_of_plots=number_of_plots, fig_nr=fig_nr, df_subs=df_subs)
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 2.3                                                    #
            # *********************************************************************************************************
            # *********************************************************************************************************
            #   END 2. Total Number of Projects and Total Capital Demand per Spatial Feature (All Years)              #
            # *********************************************************************************************************
        elif heading_no == 3:  # Capital Demand All Years
            headings3 = {"1": f"Total Capital Demand per {SpatialFeatureChoice_Text_Single} - All Years"}
            paragraphs8b = {"1 n": f"There are different funding sources from which capital is requested. For the"
                                   f" financial baseline that was selected for this report query "
                                   f"({baseline_dict['Name']}), the breakdown of capital demand for {entity_choice}"
                                   f" per funding source is shown in Figure {heading_no}.{fig_nr} below.",
                            "2 cp": f"Figure {heading_no}.{fig_nr}: Capital Demand Per Funding Source"}
            paragraphs8c = {"1 n": f"For the financial baseline that was selected for this report query "
                                   f"({baseline_dict['Name']}), the breakdown of capital demand for {entity_choice}"
                                   f" per department is shown in Figure {heading_no}.{fig_nr+1} below."
                            }
            paragraphs8d = {"1 cp": f"Figure {heading_no}.{fig_nr + 1}: Capital Demand Per {entity_choice} Department"}
            paragraphs9 = {
                "1 n": f"The following figures provide a histogram representation of the capital demand per"
                       f" {SpatialFeatureChoice_Text_Single}. A colour breakdown is provided of each year's sub-total"
                       f" that contribute towards the overall total capital demand of each "
                       f"{SpatialFeatureChoice_Text_Single}. Note that the capital demand beyond the first three years"
                       f" representing the Medium Term Expenditure Framework (MTREF - years {mtref_years_text}) are"
                       f" generally more sparsely and less diligently populated, therefore less reliance should be "
                       f"vested on these outer years when appraising the overall capital demand. The reason for this "
                       f"is usually because of National Treasury's requirement for annual budget submissions spanning"
                       f" only over the MTREF (i.e. the next three years)."}
            # *********************************************************************************************************
            #  3. Total Capital Demand per Spatial Feature - All Years                                                #
            # *********************************************************************************************************
            head_formatter(headings3)
            par_formatter(paragraphs8b)
            # *********************************************************************************************************
            #                                       Add Figure 3.1                                                    #
            # *********************************************************************************************************
            fig[fig_nr] = px.pie(df_CapexBudgetDemandCatalogue2, values='CapExDemand', names="FundingSourceName")
            fig[fig_nr].update_traces(marker=dict(colors=colors))
            fig[fig_nr].update_layout(legend=dict(yanchor="bottom", y=-0.3, xanchor="right", x=0))
            fig[fig_nr].update_layout(uniformtext_minsize=10, uniformtext_mode='show')
            fig[fig_nr].write_image(f"./static/images/fig{fig_nr}.png")
            document.add_picture(f"./static/images/fig{fig_nr}.png", width=Cm(18))
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 3.1                                                    #
            # *********************************************************************************************************
            par_formatter(paragraphs8c)
            document.add_page_break()
            par_formatter(paragraphs8d)
            # *********************************************************************************************************
            #                                       Add Figure 3.2                                                    #
            # *********************************************************************************************************
            fig[fig_nr] = px.pie(df_CapexBudgetDemandCatalogue3, values='CapExDemand', names="DepartmentName")
            fig[fig_nr].update_traces(marker=dict(colors=colors))
            fig[fig_nr].update_layout(legend=dict(yanchor="bottom", y=-0.3, xanchor="right", x=0))
            fig[fig_nr].update_layout(uniformtext_minsize=10, uniformtext_mode='show')
            fig[fig_nr].write_image(f"./static/images/fig{fig_nr}.png")
            document.add_picture(f"./static/images/fig{fig_nr}.png", width=Cm(18))
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 3.2                                                    #
            # *********************************************************************************************************
            par_formatter(paragraphs9)
            # *********************************************************************************************************
            #                                       Add Figure 3.3                                                    #
            # *********************************************************************************************************
            # Split and sort the main data frame
            df_subs = return_frames(df_Master=df_EntireSet, feature_qty=chosen_feature_qty, block_limit=40,
                                    sort_column='Capital All Years')
            number_of_plots = len(df_subs)
            # Make a histograme of the total capital demand per spatial feature stacked up with all years
            build_bar_plots_2(number_of_plots=number_of_plots, fig_nr=fig_nr, df_subs=df_subs)
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 3.3                                                    #
            # *********************************************************************************************************
            # *********************************************************************************************************
            #  END 3. Total Capital Demand per Spatial Feature - All Years                                            #
            # *********************************************************************************************************
        elif heading_no == 4:
            headings4 = {"1": f"Total Capital Demand per {SpatialFeatureChoice_Text_Single} - MTREF only"}
            paragraphs10 = {"1 n": f"Table {heading_no}.{tbl_nr} below provides a concise summary of key statistical "
                                   f"insights regarding the total MTREF capital demand per "
                                   f"{SpatialFeatureChoice_Text_Single}.",
                            "2 cp": f"Table {heading_no}.{tbl_nr}: {SpatialFeatureChoice_Text} Analysis - MTREF Capital"
                                    f" Demand (R)"}
            paragraphs11 = {"1 n": f"\nThe following figures provide a histogram representation of the capital demand"
                                   f" per {SpatialFeatureChoice_Text_Single}. A colour breakdown is provided of each "
                                   f"year's sub-total that contribute towards the overall MTREF total capital demand "
                                   f"of each {SpatialFeatureChoice_Text_Single}."}
            # *********************************************************************************************************
            #  4. Total Capital Demand per Spatial Feature - MTREF                                                    #
            # *********************************************************************************************************
            head_formatter(headings4)
            par_formatter(paragraphs10)
            # *********************************************************************************************************
            #                                       Add Table 4.1                                                     #
            # *********************************************************************************************************
            table = document.add_table(rows=1, cols=4, style='List Table 4')
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Description'
            heading_cells[1].text = 'Info (1)'
            heading_cells[2].text = 'Info (2)'
            heading_cells[3].text = 'Info (3)'
            cells = table.add_row().cells
            cells[0].text = f"MTREF Capital demand in all {SpatialFeatureChoice_Text} - Overview"
            cells[1].text = f"MTREF capital demand in all {SpatialFeatureChoice_Text}:\n{sum_cost_MTREF}"
            cells[
                2].text = f"The average MTREF capital demand per {SpatialFeatureChoice_Text_Single}:" \
                          f"\n{average_cost_MTREF}"
            cells[3].text = f"The 75th percentile of MTREF capital demand for {SpatialFeatureChoice_Text}:" \
                            f"\n{seventy_fifth_cost_MTREF}"
            cells = table.add_row().cells
            cells[0].text = f"The highest MTREF capital demand"
            cells[1].text = f"In {SpatialFeatureChoice_Text_Single}:\n{maximum_capital_feature_MTREF}"
            cells[2].text = f"MTREF capital Demand in {maximum_capital_feature_MTREF}:\n{maximum_capital_MTREF}"
            cells[
                3].text = f"Percentage of total MRTEF capital demand in {maximum_capital_feature_MTREF}:" \
                          f"\n{max_cost_perc_of_total_MTREF}"
            cells = table.add_row().cells
            cells[0].text = f"The lowest MTREF capital demand "
            cells[1].text = f"In {SpatialFeatureChoice_Text_Single}:\n{minimum_capital_feature_MTREF}"
            cells[2].text = f"MTREF capital Demand in {minimum_capital_feature_MTREF}:\n{minimum_capital_MTREF}"
            cells[
                3].text = f"Percentage of total MTREF capital demand in {minimum_capital_feature_MTREF}:" \
                          f"\n{min_cost_perc_of_total_MTREF}"
            tbl_nr += 1
            # *********************************************************************************************************
            #                                       End Table 4.1                                                     #
            # *********************************************************************************************************
            par_formatter(paragraphs11)
            # *********************************************************************************************************
            #                                       Add Figure 4.1                                                    #
            # *********************************************************************************************************
            # Split and sort the main data frame
            df_subs = return_frames(df_Master=df_EntireSet, feature_qty=chosen_feature_qty, block_limit=40,
                                    sort_column='Capital MTREF')
            number_of_plots = len(df_subs)
            # Make a histograme of the total capital demand per spatial feature stacked up with all years
            build_bar_plots_3(number_of_plots=number_of_plots, fig_nr=fig_nr, df_subs=df_subs)
            fig_nr += 1
            # *********************************************************************************************************
            #                                       End Figure 4.1                                                    #
            # *********************************************************************************************************
            # *********************************************************************************************************
            #  END 4. Total Capital Demand per Spatial Feature - MTREF                                                #
            # *********************************************************************************************************
        elif heading_no == 5:  # Summary per spatial feature
            headings5 = {"1": f"Summary per {SpatialFeatureChoice_Text_Single}"}
            # *********************************************************************************************************
            #  5. Summary per Spatial Feature                                                                         #
            # *********************************************************************************************************
            head_formatter(headings5)
            # Sort the list of features in logical order for printing in the report
            feature_str_list = re.findall(r'[A-Za-z-]+|\d+',
                                          list_of_features[0])  # Use the 1st feature to create the variable
            if len(feature_str_list) == 1 and feature_str_list[0].isnumeric():
                list_of_features.sort(key=lambda x: int(x))
            else:
                str_coord = len(feature_str_list[0])
                try:
                    list_of_features.sort(key=lambda x: int(x[str_coord:]))
                except ValueError:  # There are no numeric values to sort
                    list_of_features.sort()
            # Build the heading level 2 strings as well as the table heading strings.
            for feature in list_of_features:
                # This is to attempt to make a meaningful sub-heading if "feature" is just a number
                feature_str_list = re.findall(r'[A-Za-z-]+|\d+', feature)
                if len(feature_str_list) == 1 and feature_str_list[0].isnumeric():  # Just a number
                    from_head_str = re.findall(r'[A-Za-z-]+|\d+', SpatialFeatureChoice_Text_Single)[0]
                    if from_head_str[-1].lower() == 's':
                        from_head_str = from_head_str[:-1]
                    # Heading Level2
                    headings5b = {"2": f"{from_head_str} {feature}"}
                    paragraphs12 = {"1 cp": f"Table {heading_no}.{tbl_nr}: Summary - {from_head_str} {feature}"}
                else:
                    # Heading Level2
                    headings5b = {"2": f"{feature}"}
                    paragraphs12 = {"1 cp": f"Table {heading_no}.{tbl_nr}: Summary - {feature}"}
                # ****************************************************************************************************
                #  5.x Sub-Heading (For Each) Individual Spatial Feature                                             #
                # ****************************************************************************************************
                head_formatter(headings5b)
                # ****************************************************************************************************
                #  End 5.x Sub-Heading (For Each) Individual Spatial Feature                                         #
                # ****************************************************************************************************
                # ****************************************************************************************************
                #                         Table 5.x Summary Spatial Feature (For Each)                               #
                # ****************************************************************************************************
                par_formatter(paragraphs12)
                table = document.add_table(rows=1, cols=2, style='List Table 4')
                heading_cells = table.rows[0].cells
                heading_cells[0].text = 'Item'
                heading_cells[1].text = 'Description'
                cells = table.add_row().cells
                cells[0].text = f"Number of Projects"
                cells[1].text = f""
                tbl_nr += 1
                # ****************************************************************************************************
                #                       End Table 5.x Summary Spatial Feature (For Each)                             #
                # ****************************************************************************************************
                # ****************************************************************************************************
                #  END 5. Summary per Spatial Feature                                                                #
                # ****************************************************************************************************
        # *********************************************************************************************************
        #                                       END OF REPORT CONTENT                                             #
        # *********************************************************************************************************


    # This provides the location of the template word file
    document = Document(docx=f"./ReportFactory/Master.docx")


    # *********************************************************************************************************
    #                                       Add content - general background info                             #
    # *********************************************************************************************************
    heading_no = 0
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #                                       End - general background info                                     #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #                                       1. Introduction                                                   #
    # *********************************************************************************************************
    heading_no = 1
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #                                       End - 1. Introduction                                             #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #  2.	Total Number of Projects and Total Capital Demand per Spatial Feature (All Years)                 #
    # *********************************************************************************************************
    heading_no = 2
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #  End 2. Total Number of Projects and Total Capital Demand per Spatial Feature (All Years)               #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #  3. Total Capital Demand per Spatial Feature - All Years                                                #
    # *********************************************************************************************************
    heading_no = 3
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #  End 3. Total Capital Demand per Spatial Feature - All Years                                            #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #  4. Total Capital Demand per Spatial Feature - MTREF                                                    #
    # *********************************************************************************************************
    heading_no = 4
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #  End 4. Total Capital Demand per Spatial Feature - MTREF                                                #
    # *********************************************************************************************************


    # *********************************************************************************************************
    #  5. Summary per Spatial Feature                                                                         #
    # *********************************************************************************************************
    heading_no = 5
    fig_nr = 1
    tbl_nr = 1
    create_content(heading_no, fig_nr, tbl_nr)
    # *********************************************************************************************************
    #  End 5. Summary per Spatial Feature                                                                     #
    # *********************************************************************************************************

    full_path = f"./DOWNLOAD_FOLDER/{word_file_name}"
    document.save(full_path)

    return full_path

