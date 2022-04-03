# import os
import plotly.express as px
from docx import Document
from docx.shared import Cm
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime


def create_worddoc(var_dict, baseline_dict, df_project_cat, df_intersects2, df_subset1, df_subset2, df_subset3,
                   df_subset4, number_of_plots, df_EntireSet):
    """
    This module creates a spatial report in a MSWord File
    """
    global teller

    def par_formatter(paragraphs):
        """
        This function within a function formats paragraphs using normal, bold or italic.
        """
        paragraph = document.add_paragraph()
        for key, value in paragraphs.items():
            if " n" in key:  # Previous formatting
                paragraph.add_run(value)
            elif "al" in key: # Normal
                document.add_paragraph(value, style="Normal")
            elif " b" in key:  # Bold
                paragraph.add_run(value).bold = True
            elif " i" in key:  # Italic
                paragraph.add_run(value).italic = True
            elif " u" in key:  # Underline
                paragraph.add_run(value).underline = True
            elif "em" in key:  # Style: Emphasis
                paragraph.add_run(value).style = 'Emphasis'
            elif "tl" in key:  # Style: Title
                document.add_paragraph(value, style="Title")
            elif "st" in key:  # Style: SubTitle
                try:
                    document.add_paragraph(value, style="Subtitle")
                except KeyError:
                    document.add_paragraph(value, style="Strong")
            elif 'se' in key:
                document.add_paragraph(value, style="Section")
            elif "lb" in key:  # Style: List Bullet
                try:
                    document.add_paragraph(value, style="List Bullet")
                except KeyError:
                    document.add_paragraph(value, style="Bulleted list")
            elif "xx" in key:
                document.add_paragraph(value, style="Bulleted list")
            elif "cp" in key:  # Style: Caption
                document.add_paragraph(value, style="Caption")


    def head_formatter(headings):
        """
         This function within a function formats headings using the heading levels as input.
        """
        for key, value in headings.items():
            match key:
                case "0":
                    document.add_heading(value, level=0)
                case "1":
                    document.add_heading(value, level=1)
                case "2":
                    document.add_heading(value, level=2)
                case "3":
                    document.add_heading(value, level=3)
                case "4":
                    document.add_heading(value, level=4)


    # Put an R in front of the currency and format the amount
    def format_budget(n):
        """
        This function formats a number entry into a currency by putting an 'R' in front and delimiting with commas
        """
        currency = "R{:,}".format(int(n))
        return currency

    # Put an % in and * 100
    def format_percent(n):
        """
        This function formats a number entry into a percentage by putting an '%' in the back and delimiting with 2
        decimals
        """
        n = float(n)
        n = n * 100
        percentage = "{:.2f}%".format(n)
        return percentage


    # Unpack the variables that were passed to the function in a dictionary called var_dict
    sys_username = var_dict['username']
    url_choice = var_dict['url_choice']
    entity_choice = var_dict['entity_choice']
    SpatialFeatureChoice = var_dict['SpatialFeatureChoice']
    # Add an "s" for plural if it is not there in the word for spatial features. It reads better in the report.
    if SpatialFeatureChoice[-1].lower() != "s":
        SpatialFeatureChoice = f"{SpatialFeatureChoice}s"
    # This is required for the sake of the graphs axis titles to read right, not necessary but it just reads better
    if SpatialFeatureChoice[-1].lower()=='s':
        x_axis_values = SpatialFeatureChoice[:-1]
    else:
        x_axis_values = SpatialFeatureChoice
    Layer_List = var_dict['Layer_List']
    total_datapoints = var_dict['total_datapoints']
    intersecting = var_dict['intersecting']
    no_intersects = var_dict['no_intersects']
    perc_no_intersects = (no_intersects / total_datapoints * 100)
    chosen_feature_qty = var_dict['chosen_feature_qty']

    # Set up other vars
    timenow = datetime.today().strftime('%Y-%m-%d-%H_%M')  # Stamp the date and time
    datenow = datetime.today().strftime('%Y-%m-%d')  # Stamp the date
    urlname = url_choice.split(sep=".", maxsplit=10)[1].capitalize()
    word_file_name = f"{urlname}_SpatialRep_{timenow}.docx"

    df_EntireSet_Temp = df_EntireSet
    print(df_EntireSet.head())

    # Create variables about the projects
    maximum_projects = df_EntireSet[f'Projects per {x_axis_values}'].max()
    maximum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {x_axis_values}'] == df_EntireSet[
        f'Projects per {x_axis_values}'].max(), f'{SpatialFeatureChoice}'].iloc[0]
    minimum_projects = df_EntireSet[f'Projects per {x_axis_values}'].min()
    minimum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {x_axis_values}'] == df_EntireSet[
        f'Projects per {x_axis_values}'].min(), f'{SpatialFeatureChoice}'].iloc[0]
    average_projects = df_EntireSet[f'Projects per {x_axis_values}'].mean()
    seventy_fifth_projects = df_EntireSet[f'Projects per {x_axis_values}'].quantile(q=0.75)
    sum_projects = df_EntireSet[f'Projects per {x_axis_values}'].sum()
    # Get the top 5 with highst number of projects
    df_EntireSet_Temp.sort_values(f'Projects per {x_axis_values}', inplace=True, ascending=True)
    sum_top_five_projects = df_EntireSet_Temp[f'Projects per {x_axis_values}'].tail().sum()
    sum_top_five_projects_perc_of_total = format_percent(sum_top_five_projects/sum_projects)

    """
    # Create variables about the costs
    maximum_cost = format_budget(df_EntireSet['Capital Demand'].max())
    maximum_cost_feature = \
    df_EntireSet.loc[df_EntireSet['Capital Demand'] == df_EntireSet['Capital Demand'].max(), 'Capital Demand'].iloc[0]
    minimum_cost = format_budget(df_EntireSet['Capital Demand'].min())
    minimum_cost_feature = \
        df_EntireSet.loc[df_EntireSet['Capital Demand'] == df_EntireSet['Capital Demand'].min(), 'Capital Demand'].iloc[
            0]
    average_cost = format_budget(df_EntireSet['Capital Demand'].mean())
    seventy_fifth_cost = format_budget(df_EntireSet['Capital Demand'].quantile(q=0.75))
    sum_cost = format_budget(df_EntireSet['Capital Demand'].sum())
    max_cost_perc_of_total = format_percent(maximum_cost/sum_cost)
    df_EntireSet_Temp.sort_values('Capital Demand', inplace=True, ascending=True)
    sum_top_five_cost = format_budget(df_EntireSet_Temp['Capital Demand'].tail().sum())
    sum_top_five_cost_perc_of_total = format_percent(sum_top_five_cost/sum_cost)
    """





    # This provides the location of the template word file
    document = Document(docx=f"./DOWNLOAD_FOLDER/Master.docx")

    paragraphs00 = {}
    paragraphs00['1 tl'] = f"Spatial Query Report: {SpatialFeatureChoice}"

    # Put together the dictionary
    paragraphs0 = {}
    paragraphs0['1 n'] = "This document was created "
    paragraphs0['2 i'] = f"{timenow} "
    paragraphs0['3 n'] = "from the "
    paragraphs0['4 b'] = f"{entity_choice} "
    paragraphs0['5 n'] = "CP3 system by the following user: "
    paragraphs0['6 i'] = f"'{sys_username}'. "

    paragraphs1 = {}
    paragraphs1['1 n'] = "The baseline "
    paragraphs1['2 b'] = f"'{baseline_dict['Name']}'"
    paragraphs1['3 n'] = " with description: "
    paragraphs1['4 b'] = f"'{baseline_dict['Description']}'"
    paragraphs1['5 n'] = " was used. "
    paragraphs1['6 n'] = "The spatial feature that was selected for the purpose of this report was: "
    paragraphs1['7 b'] = f"'{SpatialFeatureChoice}'"
    paragraphs1['8 n'] = "."

    paragraphs2 = {}
    paragraphs2['1 n'] = f"This is a spatial feature report for the "
    paragraphs2['2 i'] = f"{entity_choice}."
    paragraphs2['3 n'] = f" The data contained in this report was sourced directly from the municipality's live CP3 " \
                         f"system. The information in this report therefore reflects the data as it was on the "
    paragraphs2['4 i'] = f"{datenow}"
    paragraphs2['5 n'] = f" when this report was requested. Any subsequent updates to the data contained the CP3 System" \
                         f" related to the applicable baseline from which this reports was drawn, would therefore not " \
                         f"reflect in this report. \n\nFor the "
    paragraphs2['6 i'] = f"{entity_choice}"
    paragraphs2['7 i'] = f", the following spatial features are available for the purpose of developing "
    paragraphs2['8 u'] = f"'Spatial Query Reports'"
    paragraphs2['9 n'] = f" (similar to this report):"
    teller = 10
    for feature in Layer_List:
        paragraphs2[f"{teller} bl"] = feature
        teller += 1

    paragraphs3 = {}
    paragraphs3['1 n'] = "We hope you find this useful! Sincerely, "
    paragraphs3['2 i'] = "The Novus3 Team."

    # Document Main Heading
    #head_formatter(headings0)
    par_formatter(paragraphs00)
    # Pass the dictionaries to the interpreter
    par_formatter(paragraphs0)
    par_formatter(paragraphs1)
    par_formatter(paragraphs2)
    document.add_picture(f"./static/images/CP3logo.png", width=Cm(3))
    par_formatter(paragraphs3)

    # Add a page break
    document.add_page_break()

    headings1 = {}
    headings1['1'] = "Introduction"

    # Put together the dictionary
    paragraphs4 = {}

    paragraphs4['1 n'] = f"The baseline that was queried for this report contains "
    paragraphs4['2 i'] = f"{total_datapoints}"
    paragraphs4['3 n'] = f" projects of which a total of "
    paragraphs4['4 i'] = f"{no_intersects} ({'{0:.3g}'.format(perc_no_intersects)}%)"
    paragraphs4['5 n'] = f" projects do not have any recorded intersection (overlap) with any of the following " \
                         f"spatial features:"
    teller = 6
    for feature in Layer_List:
        paragraphs4[f"{teller} xx"] = f"{feature}"
        teller += 1
    paragraphs4[f"{teller} al"] = f"\nThe probable reasons for the {no_intersects} non intersecting projects reported " \
                                  f"are that:"
    teller += 1
    paragraphs4[f"{teller} xx"] = f"these projects simply do not have a location associated with them yet (most " \
                                  f"frequently the case with 'no-intersect' projects) or;"
    teller += 1
    paragraphs4[f"{teller} xx"] = f"the location that was captured for some of these projects, may have been in " \
                                  f"the wrong place (e.g. outside the boundaries of {urlname} - this is very " \
                                  f"rarely the reason)."
    teller += 1
    paragraphs4[f"{teller} al"] = f"\nIt is important to take note of the {'{0:.3g}'.format(perc_no_intersects)}%" \
                                  f" projects that do not intersect with any spatial feature when " \
                                  f"appraising this report because a similar proportion of non-intersecting " \
                                  f"projects may be present within the specific geographic feature queried."
    if perc_no_intersects <= 5:
        teller += 1
        paragraphs4[f"{teller} al"] = f"The {'{0:.3g}'.format(perc_no_intersects)}% of non-intersecting projects is a " \
                                      f"relative low percentage and therefore not much cause of concern for this " \
                                      f"particular report."

    teller += 1
    paragraphs4[f"{teller} cp"] = "Figure 1: Intersecting vs Non-Intersecting Projects"

    head_formatter(headings1)
    par_formatter(paragraphs4)

    colors = ['red', 'green']
    fig1 = px.pie(df_intersects2, values='Projects', names='Intersects')
    fig1.update_traces(marker=dict(colors=colors, line=dict(color='#000000', width=2)))
    fig1.write_image("./static/images/fig1.png")
    document.add_picture(f"./static/images/fig1.png", width=Cm(15))

    # Add a page break
    # document.add_page_break()

    headings2 = {}
    headings2['1'] = f"{SpatialFeatureChoice} Analysis"

    paragraphs5 = {}
    paragraphs5['1 n'] = f"There is a total number of "
    paragraphs5['2 b'] = f"{chosen_feature_qty}"
    paragraphs5['3 i'] = f" {SpatialFeatureChoice}."
    paragraphs5['4 n'] = f" Each of the "
    paragraphs5['5 i'] = f"{SpatialFeatureChoice}"
    paragraphs5['6 n'] = f" has two important perspectives namely:"
    paragraphs5['7 xx'] = f"the number of  projects within each geographic area and;"
    paragraphs5['8 xx'] = f"the total capital demand per area."
    paragraphs5['9 al'] = f"\nThe following information relates to the number of projects in {SpatialFeatureChoice}:"
    paragraphs5['10 xx'] = f"{maximum_projects_feature} has the highest number of projects: {maximum_projects}" \
                           f" projects;"
    paragraphs5['11 xx'] = f"{minimum_projects_feature} has the lowest number of projects: {minimum_projects} projects;"
    paragraphs5[
        '12 xx'] = f"The average number of projects per {SpatialFeatureChoice} is {'{0:.3g}'.format(average_projects)} " \
                   f"projects;"
    paragraphs5['13 xx'] = f"The 75th percentile of projects per {SpatialFeatureChoice} is {seventy_fifth_projects} " \
                           f"projects;"
    paragraphs5['14 xx'] = f"The total number of projects in all {SpatialFeatureChoice} is {sum_projects} projects;"
    paragraphs5['15 xx'] = f"The 5 {SpatialFeatureChoice} with the most projects have a combined total of " \
                           f"{sum_top_five_projects} projects. This accounts for {sum_top_five_projects_perc_of_total}" \
                           f" of the total number of projects."

    """
    # Create variables about the projects
    maximum_projects = df_EntireSet[f'Projects per {x_axis_values}'].max()
    maximum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {x_axis_values}'] == df_EntireSet[
        'Projects per {x_axis_values}'].max(), '{x_axis_values}'].iloc[0]
    minimum_projects = df_EntireSet[f'Projects per {x_axis_values}'].min()
    minimum_projects_feature = df_EntireSet.loc[df_EntireSet[f'Projects per {x_axis_values}'] == df_EntireSet[
        'Projects per {x_axis_values}'].min(), '{x_axis_values}'].iloc[0]
    average_projects = df_EntireSet[f'Projects per {x_axis_values}'].mean()
    seventy_fifth_projects = df_EntireSet[f'Projects per {x_axis_values}'].quantile(q=0.75)
    sum_projects = df_EntireSet[f'Projects per {x_axis_values}'].sum()
    # Get the top 5 with highst number of projects
    df_EntireSet_Temp.sort_values(f'Projects per {x_axis_values}', inplace=True, ascending=True)
    sum_top_five_projects = df_EntireSet_Temp[f'Projects per {x_axis_values}'].tail().sum()
    sum_top_five_projects_perc_of_total = format_percent(sum_top_five_projects/sum_projects)
   
    # Create variables about the costs
    maximum_cost = format_budget(df_EntireSet['Capital Demand'].max())
    maximum_cost_feature = \
    df_EntireSet.loc[df_EntireSet['Capital Demand'] == df_EntireSet['Capital Demand'].max(), 'Capital Demand'].iloc[0]
    minimum_cost = format_budget(df_EntireSet['Capital Demand'].min())
    minimum_cost_feature = \
        df_EntireSet.loc[df_EntireSet['Capital Demand'] == df_EntireSet['Capital Demand'].min(), 'Capital Demand'].iloc[
            0]
    average_cost = format_budget(df_EntireSet['Capital Demand'].mean())
    seventy_fifth_cost = format_budget(df_EntireSet['Capital Demand'].quantile(q=0.75))
    sum_cost = format_budget(df_EntireSet['Capital Demand'].sum())
    max_cost_perc_of_total = format_percent(maximum_cost/sum_cost)
    df_EntireSet_Temp.sort_values('Capital Demand', inplace=True, ascending=True)
    sum_top_five_cost = format_budget(df_EntireSet_Temp['Capital Demand'].tail().sum())
    sum_top_five_cost_perc_of_total = format_percent(sum_top_five_cost/sum_cost)
    """




    head_formatter(headings2)
    par_formatter(paragraphs5)

    paragraphs6 ={}
    paragraphs7 = {}
    paragraphs8 = {}
    paragraphs9 = {}


    match number_of_plots:
        case 1:
                fig2_1 = px.bar(df_subset1, x=f'Projects per {x_axis_values}', y= SpatialFeatureChoice,
                            color='Capital Demand',
                            height=900, orientation='h')
                # Sort images to follow each other
                fig2_1.update_yaxes(categoryorder='total descending')
                # Write images to png
                fig2_1.write_image("./static/images/fig2_1.png")
                # Add a page break
                document.add_page_break()
                paragraphs6['1 cp'] = f"Figure 2: Projects and Capital Demand per {SpatialFeatureChoice}"
                par_formatter(paragraphs6)
                document.add_picture(f"./static/images/fig2_1.png", width=Cm(15))
        case 2:
                fig2_1 = px.bar(df_subset1, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                            color='Capital Demand',
                            height=900, orientation='h')
                fig2_2 = px.bar(df_subset2, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900, orientation='h')
                # Sort images to follow each other
                fig2_1.update_yaxes(categoryorder='total descending')
                fig2_2.update_yaxes(categoryorder='total descending')
                # Write images to png
                fig2_1.write_image("./static/images/fig2_1.png")
                fig2_2.write_image("./static/images/fig2_2.png")
                # Add a page break
                document.add_page_break()
                paragraphs6['1 cp'] = f"Figure 2.1: Projects and Capital Demand per {SpatialFeatureChoice} (1/2)"
                par_formatter(paragraphs6)
                document.add_picture(f"./static/images/fig2_1.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs7['1 cp'] = f"Figure 2.2: Projects and Capital Demand per {SpatialFeatureChoice} (2/2)"
                par_formatter(paragraphs7)
                document.add_picture(f"./static/images/fig2_2.png", width=Cm(15))
        case 3:
                fig2_1 = px.bar(df_subset1, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                            color='Capital Demand',
                            height=900, orientation='h')
                fig2_2 = px.bar(df_subset2, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900, orientation='h')
                fig2_3 = px.bar(df_subset3, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900, orientation='h')
                # Sort images to follow each other
                fig2_1.update_yaxes(categoryorder='total descending')
                fig2_2.update_yaxes(categoryorder='total descending')
                fig2_3.update_yaxes(categoryorder='total descending')
                # Write images to png
                fig2_1.write_image("./static/images/fig2_1.png")
                fig2_2.write_image("./static/images/fig2_2.png")
                fig2_3.write_image("./static/images/fig2_3.png")
                # Add a page break
                document.add_page_break()
                paragraphs6['1 cp'] = f"Figure 2.1: Projects and Capital Demand per {SpatialFeatureChoice} (1/3)"
                par_formatter(paragraphs6)
                document.add_picture(f"./static/images/fig2_1.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs7['1 cp'] = f"Figure 2.2: Projects and Capital Demand per {SpatialFeatureChoice} (2/3)"
                par_formatter(paragraphs7)
                document.add_picture(f"./static/images/fig2_2.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs8['1 cp'] = f"Figure 2.3: Projects and Capital Demand per {SpatialFeatureChoice} (3/3)"
                par_formatter(paragraphs8)
                document.add_picture(f"./static/images/fig2_3.png", width=Cm(15))
        case 4:
                fig2_1 = px.bar(df_subset1, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                            color='Capital Demand',
                            height=900, orientation='h')
                fig2_2 = px.bar(df_subset2, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900, orientation='h')
                fig2_3 = px.bar(df_subset3, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900, orientation='h')
                fig2_4 = px.bar(df_subset4, x=f'Projects per {x_axis_values}', y=SpatialFeatureChoice,
                                color='Capital Demand',
                                height=900)
                # Sort images to follow each other
                fig2_1.update_yaxes(categoryorder='total descending')
                fig2_2.update_yaxes(categoryorder='total descending')
                fig2_3.update_yaxes(categoryorder='total descending')
                fig2_4.update_yaxes(categoryorder='total descending')
                # Write images to png
                fig2_1.write_image("./static/images/fig2_1.png")
                fig2_2.write_image("./static/images/fig2_2.png")
                fig2_3.write_image("./static/images/fig2_3.png")
                fig2_4.write_image("./static/images/fig2_4.png")
                # Add a page break
                document.add_page_break()
                paragraphs6['1 cp'] = f"Figure 2.1: Projects and Capital Demand per {SpatialFeatureChoice} (1/4)"
                par_formatter(paragraphs6)
                document.add_picture(f"./static/images/fig2_1.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs7['1 cp'] = f"Figure 2.2: Projects and Capital Demand per {SpatialFeatureChoice} (2/4)"
                par_formatter(paragraphs7)
                document.add_picture(f"./static/images/fig2_2.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs8['1 cp'] = f"Figure 2.3: Projects and Capital Demand per {SpatialFeatureChoice} (3/4)"
                par_formatter(paragraphs8)
                document.add_picture(f"./static/images/fig2_3.png", width=Cm(15))
                # Add a page break
                document.add_page_break()
                paragraphs9['1 cp'] = f"Figure 2.4: Projects and Capital Demand per {SpatialFeatureChoice} (4/4)"
                par_formatter(paragraphs9)
                document.add_picture(f"./static/images/fig2_4.png", width=Cm(15))






    """
    There is a total number of **`r chosen_feature_qty`** `r chosen_feature`. Each of the `r chosen_feature` has two 
    important perspectives namely:`r br()`
    - the number of  projects within each geographic area and;`r br()`
    - the total capital demand per area.  `r br()`
    These two perspectives are shown in the graphics below. The lengths of the bars shown depict the number of projects 
    in `r chosen_feature`. The colour of each bar indicates the amount of capital requested within each of these areas
     - the deeper red colours are indicative of higher capital demands.
    """

    # Add a Page with General Information regarding the Municipality
    # document.add_heading(heading_1, level=1)
    # document.add_heading(heading1_1, level=2)

    #table = document.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    #heading_cells = table.rows[0].cells
    #heading_cells[0].text = 'Description'
    #heading_cells[1].text = 'Info'


    full_path = f"./DOWNLOAD_FOLDER/{word_file_name}"
    document.save(full_path)

    return full_path

